import os
import datetime
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, session
from config import Config
from models import db, Role, Department, ClassSection, User, TeacherProfile, Timetable, GPSRecord, ImageRecord, AttendanceLog, CampusSetting, TimetableUploadLog
from database import init_db
import utils
import timetable_parser
import json

app = Flask(__name__)
app.config.from_object(Config)

# Initialize Database
init_db(app)

# Ensure upload directory exists
os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'classroom_photos'), exist_ok=True)

# ----------------- SESSION AUTH DECORATOR / HELPERS -----------------
def is_logged_in():
    return 'user_id' in session

def get_current_user():
    if not is_logged_in():
        return None
    return User.query.get(session['user_id'])

# ----------------- CONTROLLER ROUTES -----------------

@app.route('/')
def index():
    if is_logged_in():
        if session.get('role') == 'admin':
            return redirect(url_for('admin_dashboard'))
        return redirect(url_for('teacher_dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if is_logged_in():
        return redirect(url_for('index'))
        
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        # Match by username OR email
        user = User.query.filter(
            (db.func.lower(User.username) == db.func.lower(username)) | 
            (db.func.lower(User.email) == db.func.lower(username))
        ).first()
        
        if user:
            # Check if account requires first-time password setup
            if user.needs_password_setup:
                session['setup_password_user_id'] = user.id
                flash("First-time login detected. Please establish your credentials.", "warning")
                return redirect(url_for('setup_password'))
                
            if user.check_password(password):
                # Check teacher approval status first
                if user.role.name == 'teacher' and user.profile:
                    status = user.profile.approval_status
                    if status == 'Pending':
                        flash("Your account registration is Pending Admin approval.", "warning")
                        return render_template('login.html')
                    elif status == 'Rejected':
                        flash("Your registration request was Rejected. Contact Admin.", "danger")
                        return render_template('login.html')
                
                # Check if account has been deactivated
                if not user.is_active:
                    flash("This account has been deactivated. Contact Admin.", "danger")
                    return render_template('login.html')
                
                session['user_id'] = user.id
                session['username'] = user.username
                session['role'] = user.role.name
                if user.role.name == 'admin':
                    session['name'] = 'System Admin'
                else:
                    session['name'] = user.profile.name if user.profile else 'Faculty Member'
                
                flash(f"Welcome back, {session['name']}!", "success")
                return redirect(url_for('index'))
            else:
                flash("Invalid username or password.", "danger")
        else:
            flash("Invalid username or password.", "danger")
            
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if is_logged_in():
        return redirect(url_for('index'))
        
    departments = Department.query.order_by(Department.name).all()
    
    if request.method == 'POST':
        employee_id = request.form.get('employee_id', '').strip()
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        phone = request.form.get('phone', '').strip()
        department_id = request.form.get('department_id')
        username = request.form.get('username', '').strip()
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        # Validation checks
        if password != confirm_password:
            flash("Passwords do not match.", "danger")
            return render_template('register.html', departments=departments)
            
        if len(password) < 8:
            flash("Password must be at least 8 characters long.", "danger")
            return render_template('register.html', departments=departments)
            
        # Check uniqueness in DB
        existing_user = User.query.filter(
            (db.func.lower(User.username) == db.func.lower(username)) |
            (db.func.lower(User.email) == db.func.lower(email))
        ).first()
        if existing_user:
            flash("Username or Email already registered.", "danger")
            return render_template('register.html', departments=departments)
            
        existing_profile = TeacherProfile.query.filter_by(employee_id=employee_id).first()
        if existing_profile:
            flash("Employee ID already registered.", "danger")
            return render_template('register.html', departments=departments)
            
        # Create User & TeacherProfile
        teacher_role = Role.query.filter_by(name='teacher').first()
        new_user = User(
            username=username,
            email=email,
            role_id=teacher_role.id,
            is_active=False, # Wait until approved
            needs_password_setup=False
        )
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        
        new_profile = TeacherProfile(
            user_id=new_user.id,
            employee_id=employee_id,
            name=name,
            phone=phone,
            department_id=int(department_id) if department_id else None,
            approval_status='Pending'
        )
        db.session.add(new_profile)
        db.session.commit()
        
        flash("Registration submitted successfully! Your profile is pending Admin approval.", "success")
        return redirect(url_for('login'))
        
    return render_template('register.html', departments=departments)

@app.route('/setup_password', methods=['GET', 'POST'])
def setup_password():
    user_id = session.get('setup_password_user_id')
    if not user_id:
        flash("Unauthorized password setup attempt.", "danger")
        return redirect(url_for('login'))
        
    user = User.query.get(user_id)
    if not user or not user.needs_password_setup:
        session.pop('setup_password_user_id', None)
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        employee_id = request.form.get('employee_id', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        # Verify identity matches seeded user email and teacher profile's employee ID
        profile = TeacherProfile.query.filter_by(user_id=user.id).first()
        
        if not profile or user.email.strip().lower() != email.lower() or profile.employee_id.strip().lower() != employee_id.lower():
            flash("Verification failed: Employee ID and Email do not match our records.", "danger")
            return render_template('setup_password.html', user_id=user_id, prefilled_email=user.email)
            
        if password != confirm_password:
            flash("Passwords do not match.", "danger")
            return render_template('setup_password.html', user_id=user_id, prefilled_email=user.email)
            
        if len(password) < 8:
            flash("Password must be at least 8 characters long.", "danger")
            return render_template('setup_password.html', user_id=user_id, prefilled_email=user.email)
            
        # Complete setup
        user.set_password(password)
        user.needs_password_setup = False
        user.is_active = True # Now fully activated
        
        # Mark profile as Approved since admin set it up
        if profile:
            profile.approval_status = 'Approved'
            
        db.session.commit()
        session.pop('setup_password_user_id', None)
        
        flash("Password established successfully! You can now log in.", "success")
        return redirect(url_for('login'))
        
    return render_template('setup_password.html', user_id=user_id, prefilled_email=user.email)

@app.route('/logout')
def logout():
    session.clear()
    flash("You have successfully signed out.", "success")
    return redirect(url_for('login'))

# ----------------- TEACHER PORTAL -----------------

@app.route('/teacher/dashboard')
def teacher_dashboard():
    if not is_logged_in() or session.get('role') != 'teacher':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    user = get_current_user()
    if not user or not user.profile:
        flash("Teacher profile not found.", "danger")
        return redirect(url_for('login'))
        
    teacher = user.profile
    now = utils.get_local_now()
    current_day = now.weekday() # 0 = Monday, ..., 6 = Sunday
    
    # Fetch campus settings to get allowed window size
    settings = CampusSetting.query.first()
    start_before = settings.start_window_before_mins if settings else 5
    end_after = settings.end_window_after_mins if settings else 15
    
    # Query timetables for this teacher for today
    timetables = Timetable.query.filter_by(teacher_profile_id=teacher.id, day_of_week=current_day).order_by(Timetable.start_time).all()
    
    schedules = []
    active_class_to_mark = None
    window_start_str = ""
    window_end_str = ""
    
    selected_mark_id = request.args.get('mark_id')
    
    for tt in timetables:
        # Check if already marked for today
        log = AttendanceLog.query.filter(
            AttendanceLog.timetable_id == tt.id,
            db.func.date(AttendanceLog.timestamp) == now.date()
        ).first()
        
        state = 'upcoming'
        marked_time = ""
        
        # Calculate allowed timing window
        dt_start = datetime.datetime.combine(now.date(), tt.start_time, tzinfo=now.tzinfo)
        dt_end = datetime.datetime.combine(now.date(), tt.end_time, tzinfo=now.tzinfo)
        allowed_start = dt_start - datetime.timedelta(minutes=start_before)
        allowed_end = dt_start + datetime.timedelta(minutes=end_after)
        
        if log:
            state = 'marked'
            marked_time = log.timestamp.strftime('%I:%M %p')
        elif allowed_start <= now <= allowed_end:
            state = 'active'
            if selected_mark_id and str(tt.id) == selected_mark_id:
                active_class_to_mark = tt
                window_start_str = allowed_start.strftime('%I:%M %p')
                window_end_str = allowed_end.strftime('%I:%M %p')
        elif now > allowed_end:
            state = 'missed'
        else:
            state = 'upcoming'
            
        schedules.append({
            'id': tt.id,
            'subject': tt.subject,
            'classroom': tt.classroom,
            'start_time': tt.start_time,
            'end_time': tt.end_time,
            'state': state,
            'marked_time': marked_time,
            'window_start': allowed_start.strftime('%I:%M %p')
        })
        
    # Auto-select the first active class if none is selected
    if not active_class_to_mark:
        for item, tt in zip(schedules, timetables):
            if item['state'] == 'active':
                active_class_to_mark = tt
                # Recalculate window times
                dt_start = datetime.datetime.combine(now.date(), tt.start_time, tzinfo=now.tzinfo)
                allowed_start = dt_start - datetime.timedelta(minutes=start_before)
                allowed_end = dt_start + datetime.timedelta(minutes=end_after)
                window_start_str = allowed_start.strftime('%I:%M %p')
                window_end_str = allowed_end.strftime('%I:%M %p')
                break

    return render_template('teacher_dashboard.html', 
                           schedules=schedules,
                           active_class_to_mark=active_class_to_mark,
                           window_start_str=window_start_str,
                           window_end_str=window_end_str,
                           current_date=now.strftime('%A, %B %d, %Y'),
                           debug_mock_gps=app.config['DEBUG_MOCK_LOCATION'])

# ----------------- BACKEND VERIFICATION API -----------------

@app.route('/api/submit_attendance', methods=['POST'])
def submit_attendance():
    if not is_logged_in() or session.get('role') != 'teacher':
        return jsonify({'message': 'Access Denied: Authentication required.'}), 401
        
    user = get_current_user()
    if not user or not user.profile:
        return jsonify({'message': 'Access Denied: Faculty profile required.'}), 401
        
    teacher = user.profile
    data = request.get_json()
    
    if not data or 'image' not in data or 'latitude' not in data or 'longitude' not in data or 'timetable_id' not in data:
        return jsonify({'message': 'Bad Request: Missing required parameters.'}), 400
        
    image_base64 = data['image']
    lat = float(data['latitude'])
    lon = float(data['longitude'])
    timetable_id = int(data['timetable_id'])
    
    now = utils.get_local_now()
    
    # 1. Fetch Timetable slot
    tt = Timetable.query.get(timetable_id)
    if not tt or tt.teacher_profile_id != teacher.id:
        return jsonify({'message': 'Timetable record not found or unauthorized.'}), 404
        
    # 2. Check if already marked for today
    existing_log = AttendanceLog.query.filter(
        AttendanceLog.timetable_id == tt.id,
        db.func.date(AttendanceLog.timestamp) == now.date()
    ).first()
    if existing_log:
        return jsonify({'message': 'Attendance already marked for this class today.'}), 400
        
    # 3. Validate Time Window
    settings = CampusSetting.query.first()
    start_before = settings.start_window_before_mins if settings else 5
    end_after = settings.end_window_after_mins if settings else 15
    
    dt_start = datetime.datetime.combine(now.date(), tt.start_time, tzinfo=now.tzinfo)
    allowed_start = dt_start - datetime.timedelta(minutes=start_before)
    allowed_end = dt_start + datetime.timedelta(minutes=end_after)
    
    if not (allowed_start <= now <= allowed_end):
        return jsonify({
            'message': 'Attendance submission blocked.',
            'reason': 'Outside allowed timing window.'
        }), 400
        
    # 4. GPS Verification
    campus_lat = settings.latitude if settings else app.config['CAMPUS_LATITUDE']
    campus_lon = settings.longitude if settings else app.config['CAMPUS_LONGITUDE']
    allowed_radius = settings.radius_meters if settings else app.config['CAMPUS_RADIUS_METERS']
    
    distance = utils.calculate_distance(lat, lon, campus_lat, campus_lon)
    
    if distance > allowed_radius:
        return jsonify({
            'message': 'Attendance rejected.',
            'reason': f'Located outside allowed campus boundary. Distance: {round(distance, 1)} meters.'
        }), 400

    # 5. Save and Compress Captured Photo
    timestamp_slug = now.strftime('%Y%m%d_%H%M%S')
    filename = f"teacher_{teacher.id}_slot_{tt.id}_{timestamp_slug}.jpg"
    image_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'classroom_photos')
    
    try:
        saved_path = utils.save_compressed_image(image_base64, image_dir, filename)
        # Relative path for storage
        relative_path = os.path.join('static', 'uploads', 'classroom_photos', filename).replace('\\', '/')
    except Exception as e:
        print(f"Image Save Error: {e}")
        return jsonify({'message': 'Error saving classroom photo on server.'}), 500
        
    # 6. Prevent Reused/Old Images (dHash Analysis)
    img_hash = utils.calculate_dhash(saved_path)
    
    # Query all old image records to look for matches
    is_duplicate = False
    duplicate_log_id = None
    
    old_images = ImageRecord.query.all()
    for old_img in old_images:
        # Avoid matching with the current file itself in case of DB conflicts
        if old_img.image_path == relative_path:
            continue
        dist = utils.hamming_distance(img_hash, old_img.image_hash)
        if dist <= 2:
            is_duplicate = True
            # Find the attendance log associated with this image record
            matching_log = AttendanceLog.query.filter_by(image_record_id=old_img.id).first()
            if matching_log:
                duplicate_log_id = matching_log.id
            break
    
    status = 'Approved'
    suspicious_reason = None
    
    if is_duplicate:
        status = 'Suspicious'
        suspicious_reason = 'Duplicate Image'
        
    # 7. AI Count Verification
    student_count = utils.estimate_human_count(saved_path)
    
    # If extremely low student count (< 3) in a class, mark suspicious
    if student_count < 3 and not is_duplicate:
        status = 'Suspicious'
        suspicious_reason = 'Low Student Count'
        
    # 8. Create GPS and Image Records (Normalized Tables)
    gps_rec = GPSRecord(
        latitude=lat,
        longitude=lon,
        distance_meters=distance,
        is_within_boundary=(distance <= allowed_radius)
    )
    db.session.add(gps_rec)
    db.session.flush()
    
    img_rec = ImageRecord(
        image_path=relative_path,
        image_hash=img_hash,
        student_count=student_count
    )
    db.session.add(img_rec)
    db.session.flush()
    
    # Insert Database Attendance Log
    log = AttendanceLog(
        teacher_profile_id=teacher.id,
        timetable_id=tt.id,
        timestamp=now,
        gps_record_id=gps_rec.id,
        image_record_id=img_rec.id,
        status=status,
        suspicious_reason=suspicious_reason
    )
    
    db.session.add(log)
    db.session.commit()
    
    if status == 'Suspicious':
        return jsonify({
            'message': 'Attendance received, but marked SUSPICIOUS. Administrative review is required.',
            'reason': suspicious_reason
        }), 200
        
    return jsonify({
        'message': 'Class presence verified successfully! Attendance marked.'
    }), 200

# ----------------- ADMIN PORTAL -----------------

@app.route('/admin/dashboard')
def admin_dashboard():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    now = utils.get_local_now()
    today_start = datetime.datetime.combine(now.date(), datetime.time.min, tzinfo=now.tzinfo)
    today_end = datetime.datetime.combine(now.date(), datetime.time.max, tzinfo=now.tzinfo)
    
    # Metrics computations for Today
    scheduled_count = Timetable.query.filter_by(day_of_week=now.weekday()).count()
    completed_logs = AttendanceLog.query.filter(AttendanceLog.timestamp.between(today_start, today_end)).all()
    completed_count = len(completed_logs)
    
    compliance_rate = 0
    if scheduled_count > 0:
        compliance_rate = round((completed_count / scheduled_count) * 100)
        
    suspicious_count = AttendanceLog.query.filter(
        AttendanceLog.timestamp.between(today_start, today_end),
        AttendanceLog.status == 'Suspicious'
    ).count()
    
    # Registrations pending approval
    pending_count = TeacherProfile.query.filter_by(approval_status='Pending').count()
    
    stats = {
        'scheduled': scheduled_count,
        'completed': completed_count,
        'compliance_rate': min(compliance_rate, 100),
        'suspicious': suspicious_count,
        'pending_registrations': pending_count
    }
    
    # Recent log entries
    recent_logs = []
    for log in completed_logs[-5:]:  # Last 5 submissions
        recent_logs.append(log.to_dict())
    recent_logs.reverse()
    
    # Flagged entries for warning list
    suspicious_logs = []
    suspicious_entries = AttendanceLog.query.filter(
        AttendanceLog.timestamp.between(today_start, today_end),
        AttendanceLog.status == 'Suspicious'
    ).all()
    for log in suspicious_entries:
        suspicious_logs.append(log.to_dict())
        
    return render_template('admin_dashboard.html', 
                           stats=stats,
                           recent_logs=recent_logs,
                           suspicious_logs=suspicious_logs)

@app.route('/admin/teachers', methods=['GET', 'POST'])
def admin_teachers():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    departments = Department.query.order_by(Department.name).all()
    teacher_role = Role.query.filter_by(name='teacher').first()
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add':
            name = request.form.get('name', '').strip()
            username = request.form.get('username', '').strip()
            email = request.form.get('email', '').strip()
            phone = request.form.get('phone', '').strip()
            employee_id = request.form.get('employee_id', '').strip()
            department_id = request.form.get('department_id')
            password = request.form.get('password')
            
            # Checks
            existing_user = User.query.filter(
                (db.func.lower(User.username) == db.func.lower(username)) |
                (db.func.lower(User.email) == db.func.lower(email))
            ).first()
            if existing_user:
                flash("Username or Email already registered.", "danger")
            else:
                existing_profile = TeacherProfile.query.filter_by(employee_id=employee_id).first()
                if existing_profile:
                    flash("Employee ID already exists.", "danger")
                else:
                    # Create user
                    needs_pw = not bool(password)
                    new_user = User(
                        username=username,
                        email=email,
                        role_id=teacher_role.id,
                        is_active=not needs_pw, # If no password, inactive until setup
                        needs_password_setup=needs_pw
                    )
                    if password:
                        new_user.set_password(password)
                    db.session.add(new_user)
                    db.session.commit()
                    
                    # Create profile
                    new_profile = TeacherProfile(
                        user_id=new_user.id,
                        employee_id=employee_id,
                        name=name,
                        phone=phone,
                        department_id=int(department_id) if department_id else None,
                        approval_status='Approved' # Admin added profiles are auto-approved
                    )
                    db.session.add(new_profile)
                    db.session.commit()
                    flash(f"Faculty profile created successfully! Password Setup required: {needs_pw}", "success")
                
        elif action == 'edit':
            profile_id = int(request.form.get('profile_id'))
            profile = TeacherProfile.query.get(profile_id)
            if profile:
                profile.name = request.form.get('name', '').strip()
                profile.phone = request.form.get('phone', '').strip()
                profile.department_id = int(request.form.get('department_id')) if request.form.get('department_id') else None
                
                user = profile.user
                if user:
                    user.email = request.form.get('email', '').strip()
                    
                    new_pwd = request.form.get('password')
                    if new_pwd:
                        user.set_password(new_pwd)
                        user.needs_password_setup = False
                        user.is_active = True
                        
                db.session.commit()
                flash("Faculty profile updated successfully!", "success")
                
        elif action == 'toggle':
            profile_id = int(request.form.get('profile_id'))
            profile = TeacherProfile.query.get(profile_id)
            if profile and profile.user:
                profile.user.is_active = not profile.user.is_active
                db.session.commit()
                state_str = "activated" if profile.user.is_active else "deactivated"
                flash(f"Faculty profile {state_str}!", "success")
                
        elif action == 'reset':
            profile_id = int(request.form.get('profile_id'))
            profile = TeacherProfile.query.get(profile_id)
            if profile and profile.user:
                profile.user.needs_password_setup = True
                profile.user.password_hash = None
                profile.user.is_active = False # Deactivate until setup again
                db.session.commit()
                flash("Faculty credentials reset. They must set up a new password on their next login.", "warning")
                
        elif action == 'delete':
            profile_id = int(request.form.get('profile_id'))
            profile = TeacherProfile.query.get(profile_id)
            if profile:
                user = profile.user
                db.session.delete(profile)
                if user:
                    db.session.delete(user)
                db.session.commit()
                flash("Faculty profile and account deleted.", "success")
                
        return redirect(url_for('admin_teachers'))
        
    profiles = TeacherProfile.query.order_by(TeacherProfile.name).all()
    edit_profile = None
    edit_id = request.args.get('edit_id')
    if edit_id:
        edit_profile = TeacherProfile.query.get(int(edit_id))
        
    return render_template('admin_teachers.html', 
                           teachers=profiles, 
                           departments=departments,
                           edit_teacher=edit_profile)

@app.route('/admin/approvals', methods=['GET', 'POST'])
def admin_approvals():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        profile_id = int(request.form.get('profile_id'))
        decision = request.form.get('decision') # 'approve' or 'reject'
        
        profile = TeacherProfile.query.get(profile_id)
        if profile and profile.user:
            if decision == 'approve':
                profile.approval_status = 'Approved'
                profile.user.is_active = True
                profile.approved_by_id = session.get('user_id')
                profile.approved_at = utils.get_local_now()
                flash(f"Approved registration for {profile.name}!", "success")
            elif decision == 'reject':
                profile.approval_status = 'Rejected'
                profile.user.is_active = False
                flash(f"Rejected registration for {profile.name}.", "warning")
                
            db.session.commit()
            
        return redirect(url_for('admin_approvals'))
        
    pending = TeacherProfile.query.filter_by(approval_status='Pending').join(TeacherProfile.user).order_by(User.created_at).all()
    past_approvals = TeacherProfile.query.filter(TeacherProfile.approval_status != 'Pending').order_by(TeacherProfile.approved_at.desc()).all()
    
    return render_template('admin_approvals.html', pending=pending, past_approvals=past_approvals)

@app.route('/admin/departments', methods=['GET', 'POST'])
def admin_departments():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add':
            name = request.form.get('name', '').strip()
            code = request.form.get('code', '').strip().upper()
            
            existing = Department.query.filter(
                (Department.code == code) | (Department.name == name)
            ).first()
            if existing:
                flash("Department name or code already exists.", "danger")
            else:
                new_dept = Department(name=name, code=code)
                db.session.add(new_dept)
                db.session.commit()
                flash("Department added successfully!", "success")
                
        elif action == 'delete':
            dept_id = int(request.form.get('department_id'))
            dept = Department.query.get(dept_id)
            if dept:
                # Any linked teachers will have their department_id set to NULL due to ON DELETE SET NULL
                db.session.delete(dept)
                db.session.commit()
                flash("Department deleted successfully.", "success")
                
        return redirect(url_for('admin_departments'))
        
    departments = Department.query.order_by(Department.name).all()
    return render_template('admin_departments.html', departments=departments)

@app.route('/admin/classes', methods=['GET', 'POST'])
def admin_classes():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add':
            name = request.form.get('name', '').strip()
            
            existing = ClassSection.query.filter_by(name=name).first()
            if existing:
                flash("Class/Course name already exists.", "danger")
            else:
                new_class = ClassSection(name=name)
                db.session.add(new_class)
                db.session.commit()
                flash("Class/Course added successfully!", "success")
                
        elif action == 'delete':
            class_id = int(request.form.get('class_id'))
            cls = ClassSection.query.get(class_id)
            if cls:
                db.session.delete(cls)
                db.session.commit()
                flash("Class/Course deleted successfully.", "success")
                
        elif action == 'rename':
            class_id = int(request.form.get('class_id'))
            new_name = request.form.get('name', '').strip()
            if not new_name:
                flash("Class name cannot be empty.", "danger")
            else:
                existing = ClassSection.query.filter_by(name=new_name).first()
                if existing:
                    flash("Class/Course name already exists.", "danger")
                else:
                    cls = ClassSection.query.get(class_id)
                    if cls:
                        cls.name = new_name
                        db.session.commit()
                        flash("Class renamed successfully!", "success")
                        
        return redirect(url_for('admin_classes'))
        
    classes = ClassSection.query.order_by(ClassSection.name).all()
    return render_template('admin_classes.html', classes=classes)

def validate_timetable_entries(entries, class_id):
    messages = []
    
    # 1. Load active teachers to format names
    teachers = TeacherProfile.query.filter_by(approval_status='Approved').all()
    teacher_map = {t.id: t.name for t in teachers}
    
    # Load all classes to print warnings nicely
    classes = ClassSection.query.all()
    class_map = {c.id: c.name for c in classes}
    target_class_name = class_map.get(class_id, "Current Class")
    
    # 2. Query all existing schedules in the database for other classes
    existing_slots = Timetable.query.filter(Timetable.class_id != class_id).all()
    
    # Parse time strings to datetime.time objects for proposed entries
    parsed_proposed = []
    for idx, entry in enumerate(entries):
        day = int(entry.get('day_of_week', 0))
        period = entry.get('period', '').strip()
        subject = entry.get('subject', '').strip()
        classroom = entry.get('classroom', '').strip()
        teacher_id = entry.get('teacher_profile_id')
        teacher_id = int(teacher_id) if (teacher_id and str(teacher_id).isdigit()) else None
        teacher_name = entry.get('teacher_name', 'Unassigned')
        
        start_str = entry.get('start_time', '09:00')
        end_str = entry.get('end_time', '10:00')
        
        try:
            start_t = datetime.datetime.strptime(start_str, '%H:%M').time()
            end_t = datetime.datetime.strptime(end_str, '%H:%M').time()
        except ValueError:
            try:
                # Fallback to parse HH:MM:SS
                start_t = datetime.datetime.strptime(start_str, '%H:%M:%S').time()
                end_t = datetime.datetime.strptime(end_str, '%H:%M:%S').time()
            except ValueError:
                start_t = datetime.time(9, 0)
                end_t = datetime.time(10, 0)
                messages.append({
                    'type': 'error',
                    'row_index': idx,
                    'message': f"Slot {idx+1}: Invalid time format '{start_str}' or '{end_str}'."
                })
            
        parsed_proposed.append({
            'index': idx,
            'day_of_week': day,
            'period': period,
            'subject': subject,
            'classroom': classroom,
            'teacher_profile_id': teacher_id,
            'teacher_name': teacher_name,
            'start_time': start_t,
            'end_time': end_t
        })
        
    DAYS_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    
    # Rule 1: Invalid timings check (start >= end)
    for p in parsed_proposed:
        if p['start_time'] >= p['end_time']:
            messages.append({
                'type': 'error',
                'row_index': p['index'],
                'message': f"Slot {p['index']+1} ({DAYS_NAMES[p['day_of_week']]} {p['period']}): Start time ({p['start_time'].strftime('%H:%M')}) must be before end time ({p['end_time'].strftime('%H:%M')})."
            })
            
    def times_overlap(s1, e1, s2, e2):
        return s1 < e2 and e1 > s2
        
    for i, p1 in enumerate(parsed_proposed):
        if p1['start_time'] >= p1['end_time']:
            continue
            
        t1_id = p1['teacher_profile_id']
        room1 = p1['classroom']
        day1 = p1['day_of_week']
        start1 = p1['start_time']
        end1 = p1['end_time']
        
        # Rule 5: Missing teacher check
        if not t1_id:
            messages.append({
                'type': 'warning',
                'row_index': p1['index'],
                'message': f"Slot {p1['index']+1} ({DAYS_NAMES[day1]} {p1['period']} - {p1['subject']}): No teacher assigned (extracted as '{p1['teacher_name']}')."
            })
            
        # Check against other proposed slots
        for j in range(i + 1, len(parsed_proposed)):
            p2 = parsed_proposed[j]
            if p2['start_time'] >= p2['end_time']:
                continue
                
            if day1 == p2['day_of_week'] and times_overlap(start1, end1, p2['start_time'], p2['end_time']):
                # Rule 2b: Teacher conflict within same upload
                if t1_id and t1_id == p2['teacher_profile_id']:
                    t_name = teacher_map.get(t1_id, p1['teacher_name'])
                    messages.append({
                        'type': 'error',
                        'row_index': p1['index'],
                        'message': f"Slot {p1['index']+1} and Slot {p2['index']+1} Conflict: Teacher '{t_name}' is double-booked at overlapping times ({start1.strftime('%H:%M')}-{end1.strftime('%H:%M')} vs {p2['start_time'].strftime('%H:%M')}-{p2['end_time'].strftime('%H:%M')}) on {DAYS_NAMES[day1]}."
                    })
                # Rule 3b: Room conflict within same upload
                if room1 and room1.lower() == p2['classroom'].lower() and room1.lower() != "room" and room1.strip():
                    messages.append({
                        'type': 'error',
                        'row_index': p1['index'],
                        'message': f"Slot {p1['index']+1} and Slot {p2['index']+1} Conflict: Room '{room1}' is double-booked at overlapping times on {DAYS_NAMES[day1]}."
                    })
                # Rule 4: Class overlap within same upload
                messages.append({
                    'type': 'warning',
                    'row_index': p1['index'],
                    'message': f"Slot {p1['index']+1} and Slot {p2['index']+1} Overlap: {target_class_name} has multiple overlapping classes scheduled on {DAYS_NAMES[day1]}."
                })
                
        # Check against existing database slots of OTHER classes
        for db_slot in existing_slots:
            if day1 == db_slot.day_of_week and times_overlap(start1, end1, db_slot.start_time, db_slot.end_time):
                other_class_name = class_map.get(db_slot.class_id, "Another Class")
                # Rule 2c: Teacher Conflict with database
                if t1_id and t1_id == db_slot.teacher_profile_id:
                    t_name = teacher_map.get(t1_id, p1['teacher_name'])
                    messages.append({
                        'type': 'error',
                        'row_index': p1['index'],
                        'message': f"Slot {p1['index']+1} Conflict: Teacher '{t_name}' is already assigned to {other_class_name} ({db_slot.subject}) in room {db_slot.classroom} at overlapping time ({db_slot.start_time.strftime('%H:%M')}-{db_slot.end_time.strftime('%H:%M')}) on {DAYS_NAMES[day1]}."
                    })
                # Rule 3c: Room Conflict with database
                if room1 and room1.lower() == db_slot.classroom.lower() and room1.lower() != "room" and room1.strip():
                    messages.append({
                        'type': 'error',
                        'row_index': p1['index'],
                        'message': f"Slot {p1['index']+1} Conflict: Room '{room1}' is already booked by {other_class_name} (Teacher: {db_slot.teacher_profile.name if db_slot.teacher_profile else 'Unknown'}) at overlapping time ({db_slot.start_time.strftime('%H:%M')}-{db_slot.end_time.strftime('%H:%M')}) on {DAYS_NAMES[day1]}."
                    })
                    
    return messages

@app.route('/admin/timetable/upload', methods=['POST'])
def admin_timetable_upload():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    class_id_str = request.form.get('class_id')
    file = request.files.get('timetable_file')
    
    if not class_id_str or not file or file.filename == '':
        flash("Please select a class and select a timetable file.", "danger")
        return redirect(url_for('admin_timetable'))
        
    class_id = int(class_id_str)
    cls = ClassSection.query.get(class_id)
    if not cls:
        flash("Selected class does not exist.", "danger")
        return redirect(url_for('admin_timetable'))
        
    # Save file
    os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'timetables'), exist_ok=True)
    filename = file.filename
    safe_filename = f"class_{class_id}_{int(datetime.datetime.utcnow().timestamp())}_{filename}"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'timetables', safe_filename)
    file.save(filepath)
    
    # Get active teachers
    teachers_raw = TeacherProfile.query.filter_by(approval_status='Approved').join(TeacherProfile.user).filter(User.is_active == True).order_by(TeacherProfile.name).all()
    teachers_list = [{'id': t.id, 'name': t.name} for t in teachers_raw]
    
    ext = filename.split('.')[-1]
    
    try:
        parsed_slots = timetable_parser.parse_timetable_file(filepath, ext, teachers_list)
    except Exception as e:
        flash(f"Timetable Parse Error: {str(e)}", "danger")
        return redirect(url_for('admin_timetable'))
        
    # Validate parsed slots
    warnings = validate_timetable_entries(parsed_slots, class_id)
    
    # Save log entry
    log = TimetableUploadLog(
        class_id=class_id,
        filename=filename,
        filepath=f"static/uploads/timetables/{safe_filename}",
        status="Draft",
        import_summary=f"Parsed {len(parsed_slots)} entries",
        validation_errors=json.dumps(warnings)
    )
    db.session.add(log)
    db.session.commit()
    
    # Render interactive editor page
    return render_template('admin_timetable_review.html',
                           class_id=class_id,
                           class_name=cls.name,
                           filename=filename,
                           log_id=log.id,
                           slots=parsed_slots,
                           warnings=warnings,
                           teachers=teachers_raw)

@app.route('/admin/timetable/validate', methods=['POST'])
def admin_timetable_validate():
    if not is_logged_in() or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Access Denied'}), 403
        
    data = request.json or {}
    class_id = int(data.get('class_id', 0))
    entries = data.get('entries', [])
    
    warnings = validate_timetable_entries(entries, class_id)
    return jsonify({'success': True, 'warnings': warnings})

@app.route('/admin/timetable/save-uploaded', methods=['POST'])
def admin_timetable_save_uploaded():
    if not is_logged_in() or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Access Denied'}), 403
        
    data = request.json or {}
    class_id = int(data.get('class_id', 0))
    log_id = int(data.get('log_id', 0))
    entries = data.get('entries', [])
    
    cls = ClassSection.query.get(class_id)
    if not cls:
        return jsonify({'success': False, 'message': 'Class does not exist'}), 400
        
    log = TimetableUploadLog.query.get(log_id)
    if not log:
        return jsonify({'success': False, 'message': 'Upload log does not exist'}), 400
        
    for idx, entry in enumerate(entries):
        teacher_id = entry.get('teacher_profile_id')
        if not teacher_id or str(teacher_id).lower() == 'null' or str(teacher_id).strip() == '':
            return jsonify({
                'success': False, 
                'message': f"Slot {idx+1} ({entry.get('subject')}) is unassigned. Please map a teacher before saving."
            }), 400
            
    # Delete old slots
    Timetable.query.filter_by(class_id=class_id).delete()
    
    # Save new slots
    for entry in entries:
        start_t = datetime.datetime.strptime(entry['start_time'], '%H:%M').time()
        end_t = datetime.datetime.strptime(entry['end_time'], '%H:%M').time()
        
        tt = Timetable(
            teacher_profile_id=int(entry['teacher_profile_id']),
            class_id=class_id,
            subject=entry['subject'].strip(),
            classroom=entry['classroom'].strip(),
            day_of_week=int(entry['day_of_week']),
            period=entry['period'].strip(),
            start_time=start_t,
            end_time=end_t
        )
        db.session.add(tt)
        
    log.status = "Completed"
    log.import_summary = f"Successfully imported {len(entries)} slots"
    log.validation_errors = "[]"
    
    db.session.commit()
    
    flash(f"Timetable for class {cls.name} updated successfully! {len(entries)} periods scheduled.", "success")
    
    return jsonify({
        'success': True,
        'redirect_url': url_for('admin_timetable')
    })

@app.route('/admin/timetable', methods=['GET', 'POST'])
def admin_timetable():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add':
            profile_id = int(request.form.get('teacher_profile_id'))
            class_id = int(request.form.get('class_id'))
            subject = request.form.get('subject', '').strip()
            classroom = request.form.get('classroom', '').strip()
            day_of_week = int(request.form.get('day_of_week'))
            period = request.form.get('period', 'Period 1').strip()
            
            start_str = request.form.get('start_time')
            end_str = request.form.get('end_time')
            
            # Parse times
            start_time = datetime.datetime.strptime(start_str, '%H:%M').time()
            end_time = datetime.datetime.strptime(end_str, '%H:%M').time()
            
            tt = Timetable(
                teacher_profile_id=profile_id,
                class_id=class_id,
                subject=subject,
                classroom=classroom,
                day_of_week=day_of_week,
                period=period,
                start_time=start_time,
                end_time=end_time
            )
            db.session.add(tt)
            db.session.commit()
            flash("Schedule mapped successfully!", "success")
            
        elif action == 'delete':
            tt_id = int(request.form.get('timetable_id'))
            tt = Timetable.query.get(tt_id)
            if tt:
                db.session.delete(tt)
                db.session.commit()
                flash("Schedule mapping removed.", "success")
                
        return redirect(url_for('admin_timetable'))
        
    teachers = TeacherProfile.query.filter_by(approval_status='Approved').join(TeacherProfile.user).filter(User.is_active == True).order_by(TeacherProfile.name).all()
    classes = ClassSection.query.order_by(ClassSection.name).all()
    timetables_raw = Timetable.query.join(TeacherProfile).order_by(Timetable.day_of_week, Timetable.start_time).all()
    
    timetables = [tt.to_dict() for tt in timetables_raw]
    upload_logs = TimetableUploadLog.query.order_by(TimetableUploadLog.uploaded_at.desc()).all()
    return render_template('admin_timetable.html', 
                           teachers=teachers, 
                           classes=classes,
                           timetables=timetables,
                           upload_logs=upload_logs)

@app.route('/admin/attendance')
def admin_attendance():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    # Filters parameters
    profile_id = request.args.get('profile_id')
    status = request.args.get('status')
    date_str = request.args.get('date')
    reason = request.args.get('reason')
    
    query = AttendanceLog.query
    
    if profile_id:
        query = query.filter(AttendanceLog.teacher_profile_id == int(profile_id))
    if status:
        query = query.filter(AttendanceLog.status == status)
    if date_str:
        target_date = datetime.datetime.strptime(date_str, '%Y-%m-%d').date()
        query = query.filter(db.func.date(AttendanceLog.timestamp) == target_date)
    if reason:
        query = query.filter(AttendanceLog.suspicious_reason == reason)
        
    logs_raw = query.order_by(AttendanceLog.timestamp.desc()).all()
    logs = [log.to_dict() for log in logs_raw]
    
    teachers = TeacherProfile.query.order_by(TeacherProfile.name).all()
    
    return render_template('admin_attendance.html',
                           logs=logs,
                           teachers=teachers,
                           selected_profile_id=profile_id,
                           selected_status=status,
                           selected_date=date_str,
                           selected_reason=reason)

@app.route('/admin/verify_action', methods=['POST'])
def admin_verify_action():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    log_id = int(request.form.get('log_id'))
    status = request.form.get('status')
    notes = request.form.get('notes')
    
    log = AttendanceLog.query.get(log_id)
    if log:
        log.status = status
        log.verification_notes = notes
        db.session.commit()
        flash("Attendance verification log updated successfully.", "success")
    else:
        flash("Log record not found.", "danger")
        
    return redirect(url_for('admin_attendance'))

@app.route('/admin/settings', methods=['GET', 'POST'])
def admin_settings():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    settings = CampusSetting.query.first()
    
    if request.method == 'POST':
        campus_name = request.form.get('campus_name', '').strip()
        latitude = float(request.form.get('latitude'))
        longitude = float(request.form.get('longitude'))
        radius_meters = float(request.form.get('radius_meters'))
        start_window_before_mins = int(request.form.get('start_window_before_mins'))
        end_window_after_mins = int(request.form.get('end_window_after_mins'))
        
        if not settings:
            settings = CampusSetting()
            db.session.add(settings)
            
        settings.campus_name = campus_name
        settings.latitude = latitude
        settings.longitude = longitude
        settings.radius_meters = radius_meters
        settings.start_window_before_mins = start_window_before_mins
        settings.end_window_after_mins = end_window_after_mins
        
        db.session.commit()
        flash("Campus Rules and coordinates updated successfully!", "success")
        return redirect(url_for('admin_settings'))
        
    return render_template('admin_settings.html', settings=settings)

@app.route('/admin/reports')
def admin_reports():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    teacher_id = request.args.get('teacher_id')
    department_id = request.args.get('department_id')
    class_id = request.args.get('class_id')
    date_range = request.args.get('date_range', 'all')
    
    # Build filtered query for logs
    log_query = AttendanceLog.query.join(TeacherProfile).outerjoin(Department, TeacherProfile.department_id == Department.id).join(Timetable, AttendanceLog.timetable_id == Timetable.id)
    
    if teacher_id:
        log_query = log_query.filter(AttendanceLog.teacher_profile_id == int(teacher_id))
    if department_id:
        log_query = log_query.filter(TeacherProfile.department_id == int(department_id))
    if class_id:
        log_query = log_query.filter(Timetable.class_id == int(class_id))
        
    if date_range and date_range != 'all':
        now = utils.get_local_now()
        if date_range == 'today':
            start = datetime.datetime.combine(now.date(), datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(now.date(), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'week':
            start_of_week = now.date() - datetime.timedelta(days=now.weekday())
            start = datetime.datetime.combine(start_of_week, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(start_of_week + datetime.timedelta(days=6), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'month':
            start_of_month = now.date().replace(day=1)
            next_month = start_of_month + datetime.timedelta(days=32)
            end_of_month = next_month.replace(day=1) - datetime.timedelta(days=1)
            start = datetime.datetime.combine(start_of_month, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(end_of_month, datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
            
    filtered_logs = log_query.all()
    
    # Calculate counters based on filtered logs
    total_logs = len(filtered_logs)
    total_suspicious = sum(1 for log in filtered_logs if log.status == 'Suspicious')
    
    # Specific flag incident counts
    duplicate_images = sum(1 for log in filtered_logs if log.suspicious_reason == 'Duplicate Image')
    outside_campus = sum(1 for log in filtered_logs if log.suspicious_reason == 'Outside Campus')
    low_students = sum(1 for log in filtered_logs if log.suspicious_reason == 'Low Student Count')
    
    incidents = {
        'duplicate_images': duplicate_images,
        'outside_campus': outside_campus,
        'low_students': low_students
    }
    
    # Filtered timetable slots count
    timetable_query = Timetable.query.join(TeacherProfile)
    if teacher_id:
        timetable_query = timetable_query.filter(Timetable.teacher_profile_id == int(teacher_id))
    if department_id:
        timetable_query = timetable_query.filter(TeacherProfile.department_id == int(department_id))
    if class_id:
        timetable_query = timetable_query.filter(Timetable.class_id == int(class_id))
    total_scheduled = timetable_query.count()
    
    # Teacher statistics compiling
    teacher_query = TeacherProfile.query
    if teacher_id:
        teacher_query = teacher_query.filter(TeacherProfile.id == int(teacher_id))
    if department_id:
        teacher_query = teacher_query.filter(TeacherProfile.department_id == int(department_id))
        
    teachers_list = teacher_query.all()
    teacher_stats = []
    
    for t in teachers_list:
        t_logs_q = AttendanceLog.query.filter_by(teacher_profile_id=t.id).join(Timetable)
        if class_id:
            t_logs_q = t_logs_q.filter(Timetable.class_id == int(class_id))
            
        if date_range and date_range != 'all':
            now = utils.get_local_now()
            if date_range == 'today':
                start = datetime.datetime.combine(now.date(), datetime.time.min, tzinfo=now.tzinfo)
                end = datetime.datetime.combine(now.date(), datetime.time.max, tzinfo=now.tzinfo)
                t_logs_q = t_logs_q.filter(AttendanceLog.timestamp.between(start, end))
            elif date_range == 'week':
                start_of_week = now.date() - datetime.timedelta(days=now.weekday())
                start = datetime.datetime.combine(start_of_week, datetime.time.min, tzinfo=now.tzinfo)
                end = datetime.datetime.combine(start_of_week + datetime.timedelta(days=6), datetime.time.max, tzinfo=now.tzinfo)
                t_logs_q = t_logs_q.filter(AttendanceLog.timestamp.between(start, end))
            elif date_range == 'month':
                start_of_month = now.date().replace(day=1)
                next_month = start_of_month + datetime.timedelta(days=32)
                end_of_month = next_month.replace(day=1) - datetime.timedelta(days=1)
                start = datetime.datetime.combine(start_of_month, datetime.time.min, tzinfo=now.tzinfo)
                end = datetime.datetime.combine(end_of_month, datetime.time.max, tzinfo=now.tzinfo)
                t_logs_q = t_logs_q.filter(AttendanceLog.timestamp.between(start, end))
                
        marked_count = t_logs_q.count()
        susp_count = t_logs_q.filter(AttendanceLog.status == 'Suspicious').count()
        
        t_slots_q = Timetable.query.filter_by(teacher_profile_id=t.id)
        if class_id:
            t_slots_q = t_slots_q.filter(Timetable.class_id == int(class_id))
        slots_count = t_slots_q.count()
        
        rate = 100
        if slots_count > 0:
            if date_range == 'today':
                opps = slots_count
            elif date_range == 'week':
                opps = slots_count * 5
            elif date_range == 'month':
                opps = slots_count * 20
            else:
                opps = slots_count * 20
            rate = round((marked_count / opps) * 100)
            rate = min(rate, 100)
        else:
            rate = 0
            
        teacher_stats.append({
            'name': t.name,
            'department': t.department.name if t.department else 'N/A',
            'marked_count': marked_count,
            'suspicious_count': susp_count,
            'rate': rate
        })
        
    all_teachers = TeacherProfile.query.order_by(TeacherProfile.name).all()
    all_departments = Department.query.order_by(Department.name).all()
    all_classes = ClassSection.query.order_by(ClassSection.name).all()
    
    return render_template('admin_reports.html',
                           total_scheduled_all_time=total_scheduled,
                           total_logs_all_time=total_logs,
                           total_suspicious_all_time=total_suspicious,
                           teacher_stats=teacher_stats,
                           incidents=incidents,
                           teachers=all_teachers,
                           departments=all_departments,
                           classes=all_classes,
                           selected_teacher_id=teacher_id,
                           selected_department_id=department_id,
                           selected_class_id=class_id,
                           selected_date_range=date_range)

@app.route('/admin/reports/download/pdf')
def admin_reports_download_pdf():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    teacher_id = request.args.get('teacher_id')
    department_id = request.args.get('department_id')
    class_id = request.args.get('class_id')
    date_range = request.args.get('date_range')
    
    desc_parts = []
    log_query = AttendanceLog.query.join(TeacherProfile).outerjoin(Department, TeacherProfile.department_id == Department.id).join(Timetable, AttendanceLog.timetable_id == Timetable.id)
    
    if teacher_id and teacher_id != '':
        log_query = log_query.filter(AttendanceLog.teacher_profile_id == int(teacher_id))
        teacher = TeacherProfile.query.get(int(teacher_id))
        if teacher:
            desc_parts.append(f"Teacher: {teacher.name}")
    else:
        desc_parts.append("Teachers: All")
        
    if department_id and department_id != '':
        log_query = log_query.filter(TeacherProfile.department_id == int(department_id))
        dept = Department.query.get(int(department_id))
        if dept:
            desc_parts.append(f"Department: {dept.name}")
    else:
        desc_parts.append("Departments: All")
        
    if class_id and class_id != '':
        log_query = log_query.filter(Timetable.class_id == int(class_id))
        cls = ClassSection.query.get(int(class_id))
        if cls:
            desc_parts.append(f"Class: {cls.name}")
    else:
        desc_parts.append("Classes: All")
        
    if date_range and date_range != '' and date_range != 'all':
        now = utils.get_local_now()
        desc_parts.append(f"Date Range: {date_range.capitalize()}")
        if date_range == 'today':
            start = datetime.datetime.combine(now.date(), datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(now.date(), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'week':
            start_of_week = now.date() - datetime.timedelta(days=now.weekday())
            start = datetime.datetime.combine(start_of_week, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(start_of_week + datetime.timedelta(days=6), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'month':
            start_of_month = now.date().replace(day=1)
            next_month = start_of_month + datetime.timedelta(days=32)
            end_of_month = next_month.replace(day=1) - datetime.timedelta(days=1)
            start = datetime.datetime.combine(start_of_month, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(end_of_month, datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
    else:
        desc_parts.append("Date Range: All Time")
        
    logs_raw = log_query.order_by(AttendanceLog.timestamp.desc()).all()
    logs = [log.to_dict() for log in logs_raw]
    filters_desc = " | ".join(desc_parts)
    
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    import io
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1e1b4b'),
        spaceAfter=12
    )
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#4b5563'),
        spaceAfter=20
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#1e1b4b'),
        spaceBefore=15,
        spaceAfter=8
    )
    th_style = ParagraphStyle(
        'TableHead',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )
    td_style = ParagraphStyle(
        'TableBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#1f2937')
    )
    td_bold_style = ParagraphStyle(
        'TableBodyBold',
        parent=td_style,
        fontName='Helvetica-Bold'
    )
    
    story.append(Paragraph("VeriClass - Academic Compliance Report", title_style))
    story.append(Paragraph(f"<b>Report Generation Time (IST):</b> {utils.get_local_now().strftime('%Y-%m-%d %I:%M %p')}<br/><b>Applied Filters:</b> {filters_desc}", subtitle_style))
    
    total_logs = len(logs)
    suspicious_count = sum(1 for log in logs if log['status'] == 'Suspicious')
    approved_count = sum(1 for log in logs if log['status'] == 'Approved')
    rejected_count = sum(1 for log in logs if log['status'] == 'Rejected')
    
    kpi_data = [
        [Paragraph("<b>Metric</b>", td_bold_style), Paragraph("<b>Count</b>", td_bold_style), Paragraph("<b>Percentage</b>", td_bold_style)],
        [Paragraph("Total Submissions", td_style), Paragraph(str(total_logs), td_style), Paragraph("100%", td_style)],
        [Paragraph("Approved Submissions", td_style), Paragraph(str(approved_count), td_style), Paragraph(f"{round(approved_count/total_logs*100) if total_logs > 0 else 0}%", td_style)],
        [Paragraph("Suspicious Flags", td_style), Paragraph(str(suspicious_count), td_style), Paragraph(f"{round(suspicious_count/total_logs*100) if total_logs > 0 else 0}%", td_style)],
        [Paragraph("Rejected Submissions", td_style), Paragraph(str(rejected_count), td_style), Paragraph(f"{round(rejected_count/total_logs*100) if total_logs > 0 else 0}%", td_style)],
    ]
    
    kpi_table = Table(kpi_data, colWidths=[200, 170, 170])
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f3f4f6')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
    ]))
    
    story.append(Paragraph("KPI Summary Statistics", h2_style))
    story.append(kpi_table)
    story.append(Spacer(1, 15))
    
    story.append(Paragraph("Detailed Compliance Logs", h2_style))
    
    log_headers = ["Faculty", "Class", "Period", "Subject", "Date & Time", "Status", "Distance"]
    table_data = [[Paragraph(h, th_style) for h in log_headers]]
    
    for l in logs:
        ts = l['timestamp']
        try:
            dt = datetime.datetime.strptime(ts, '%Y-%m-%d %H:%M:%S')
            ts_formatted = dt.strftime('%m/%d %I:%M %p')
        except:
            ts_formatted = ts
            
        status_color = '#10b981' if l['status'] == 'Approved' else ('#f59e0b' if l['status'] == 'Suspicious' else '#ef4444')
        status_para = Paragraph(f"<font color='{status_color}'><b>{l['status']}</b></font>", td_bold_style)
        
        table_data.append([
            Paragraph(l['teacher_name'], td_bold_style),
            Paragraph(l['class_name'], td_style),
            Paragraph(l['period'], td_style),
            Paragraph(l['subject'], td_style),
            Paragraph(ts_formatted, td_style),
            status_para,
            Paragraph(f"{l['distance_meters']} m", td_style)
        ])
        
    logs_table = Table(table_data, colWidths=[85, 80, 65, 85, 90, 75, 60], repeatRows=1)
    logs_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e1b4b')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
    ]))
    story.append(logs_table)
    
    doc.build(story)
    buffer.seek(0)
    
    from flask import make_response
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=compliance_report.pdf'
    return response

@app.route('/admin/reports/download/docx')
def admin_reports_download_docx():
    if not is_logged_in() or session.get('role') != 'admin':
        flash("Access Denied.", "danger")
        return redirect(url_for('login'))
        
    teacher_id = request.args.get('teacher_id')
    department_id = request.args.get('department_id')
    class_id = request.args.get('class_id')
    date_range = request.args.get('date_range')
    
    desc_parts = []
    log_query = AttendanceLog.query.join(TeacherProfile).outerjoin(Department, TeacherProfile.department_id == Department.id).join(Timetable, AttendanceLog.timetable_id == Timetable.id)
    
    if teacher_id and teacher_id != '':
        log_query = log_query.filter(AttendanceLog.teacher_profile_id == int(teacher_id))
        teacher = TeacherProfile.query.get(int(teacher_id))
        if teacher:
            desc_parts.append(f"Teacher: {teacher.name}")
    else:
        desc_parts.append("Teachers: All")
        
    if department_id and department_id != '':
        log_query = log_query.filter(TeacherProfile.department_id == int(department_id))
        dept = Department.query.get(int(department_id))
        if dept:
            desc_parts.append(f"Department: {dept.name}")
    else:
        desc_parts.append("Departments: All")
        
    if class_id and class_id != '':
        log_query = log_query.filter(Timetable.class_id == int(class_id))
        cls = ClassSection.query.get(int(class_id))
        if cls:
            desc_parts.append(f"Class: {cls.name}")
    else:
        desc_parts.append("Classes: All")
        
    if date_range and date_range != '' and date_range != 'all':
        now = utils.get_local_now()
        desc_parts.append(f"Date Range: {date_range.capitalize()}")
        if date_range == 'today':
            start = datetime.datetime.combine(now.date(), datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(now.date(), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'week':
            start_of_week = now.date() - datetime.timedelta(days=now.weekday())
            start = datetime.datetime.combine(start_of_week, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(start_of_week + datetime.timedelta(days=6), datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
        elif date_range == 'month':
            start_of_month = now.date().replace(day=1)
            next_month = start_of_month + datetime.timedelta(days=32)
            end_of_month = next_month.replace(day=1) - datetime.timedelta(days=1)
            start = datetime.datetime.combine(start_of_month, datetime.time.min, tzinfo=now.tzinfo)
            end = datetime.datetime.combine(end_of_month, datetime.time.max, tzinfo=now.tzinfo)
            log_query = log_query.filter(AttendanceLog.timestamp.between(start, end))
    else:
        desc_parts.append("Date Range: All Time")
        
    logs_raw = log_query.order_by(AttendanceLog.timestamp.desc()).all()
    logs = [log.to_dict() for log in logs_raw]
    filters_desc = " | ".join(desc_parts)
    
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    import io
    
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    title = doc.add_paragraph()
    title_run = title.add_run("VeriClass - Academic Compliance Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 27, 75)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    meta = doc.add_paragraph()
    meta.add_run(f"Report Generation Time (IST): ").bold = True
    meta.add_run(f"{utils.get_local_now().strftime('%Y-%m-%d %I:%M %p')}\n")
    meta.add_run(f"Applied Filters: ").bold = True
    meta.add_run(f"{filters_desc}\n")
    for run in meta.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 85, 99)
    
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("KPI Summary Statistics")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(12)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(30, 27, 75)
    
    total_logs = len(logs)
    suspicious_count = sum(1 for log in logs if log['status'] == 'Suspicious')
    approved_count = sum(1 for log in logs if log['status'] == 'Approved')
    rejected_count = sum(1 for log in logs if log['status'] == 'Rejected')
    
    kpi_table = doc.add_table(rows=5, cols=3)
    kpi_table.style = 'Light Shading Accent 1'
    
    kpi_headers = ["Metric", "Count", "Percentage"]
    for i, head in enumerate(kpi_headers):
        kpi_table.rows[0].cells[i].text = head
        kpi_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
    kpi_rows = [
        ("Total Submissions", str(total_logs), "100%"),
        ("Approved Submissions", str(approved_count), f"{round(approved_count/total_logs*100) if total_logs > 0 else 0}%"),
        ("Suspicious Flags", str(suspicious_count), f"{round(suspicious_count/total_logs*100) if total_logs > 0 else 0}%"),
        ("Rejected Submissions", str(rejected_count), f"{round(rejected_count/total_logs*100) if total_logs > 0 else 0}%"),
    ]
    
    for row_idx, data in enumerate(kpi_rows):
        for col_idx, text in enumerate(data):
            kpi_table.rows[row_idx + 1].cells[col_idx].text = text
            
    doc.add_paragraph()
    
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Detailed Compliance Logs")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(12)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(30, 27, 75)
    
    log_headers = ["Faculty", "Class", "Period", "Subject", "Date & Time", "Status", "Distance"]
    log_table = doc.add_table(rows=1, cols=len(log_headers))
    log_table.style = 'Table Grid'
    
    hdr_cells = log_table.rows[0].cells
    for i, t_title in enumerate(log_headers):
        hdr_cells[i].text = t_title
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd {} w:fill="1E1B4B"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
    for l in logs:
        row_cells = log_table.add_row().cells
        row_cells[0].text = l['teacher_name']
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = l['class_name']
        row_cells[2].text = l['period']
        row_cells[3].text = l['subject']
        
        ts = l['timestamp']
        try:
            dt = datetime.datetime.strptime(ts, '%Y-%m-%d %H:%M:%S')
            ts_formatted = dt.strftime('%m/%d %I:%M %p')
        except:
            ts_formatted = ts
        row_cells[4].text = ts_formatted
        
        status_para = row_cells[5].paragraphs[0]
        status_run = status_para.add_run(l['status'])
        status_run.font.bold = True
        if l['status'] == 'Approved':
            status_run.font.color.rgb = RGBColor(16, 185, 129)
        elif l['status'] == 'Suspicious':
            status_run.font.color.rgb = RGBColor(245, 158, 11)
        else:
            status_run.font.color.rgb = RGBColor(239, 68, 68)
            
        row_cells[6].text = f"{l['distance_meters']} m"
        
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    from flask import make_response
    response = make_response(doc_stream.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = 'attachment; filename=compliance_report.docx'
    return response

# Start Application
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
